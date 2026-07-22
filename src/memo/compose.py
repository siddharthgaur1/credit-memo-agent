"""Credit memo composition and .docx export.

The memo is a DRAFT for a credit officer to verify and own. It presents evidence
-- figures, formulas, citations, gaps -- and deliberately issues no verdict, no
recommendation and no score. That is the honest design and the only one that
survives an audit.
"""

from __future__ import annotations

from datetime import date
from pathlib import Path

from pydantic import BaseModel, Field

from ..analysis.ratios import AnalysisResult, Completeness
from ..analysis.risk import RiskFlag
from ..checklist.validate import ChecklistReport, Status
from ..extract.schemas import ExtractedFinancials, Figure, val
from ..extract.verify import VerificationFinding

DISCLAIMER = (
    "MACHINE-GENERATED DRAFT -- NOT A CREDIT DECISION. "
    "This memo was assembled automatically from the documents submitted with this file. "
    "Every figure is cited to its source and must be verified against the source document "
    "before use. It contains no recommendation, no approval or rejection, and no credit "
    "score. The credit officer owns the assessment and the decision."
)

SUMMARY_ROWS: list[tuple[str, str]] = [
    ("Revenue", "pl.revenue"),
    ("EBITDA", "pl.ebitda"),
    ("Finance cost", "pl.finance_cost"),
    ("Profit after tax", "pl.profit_after_tax"),
    ("Net worth", "bs.net_worth"),
    ("Total debt", "bs._total_debt"),
    ("Total assets", "bs.total_assets"),
    ("Total current assets", "bs.total_current_assets"),
    ("Total current liabilities", "bs.total_current_liabilities"),
    ("Inventory", "bs.inventory"),
    ("Trade receivables", "bs.trade_receivables"),
    ("Trade payables", "bs.trade_payables"),
]


class MemoSection(BaseModel):
    heading: str
    body: str = ""
    table: list[list[str]] = Field(default_factory=list)  # first row = header


class CreditMemo(BaseModel):
    borrower_name: str
    facility_requested: str
    prepared_on: date
    backend: str
    sections: list[MemoSection]
    citations: list[str] = Field(default_factory=list)
    open_questions: list[str] = Field(default_factory=list)


def _fmt(f: Figure | float | None) -> str:
    if f is None:
        return "not found"
    value = f.value if isinstance(f, Figure) else f
    return f"{value:,.0f}"


def _summary_table(fin: ExtractedFinancials) -> list[list[str]]:
    periods = fin.periods()
    rows = [["Line item (Rs)"] + periods]
    for label, key in SUMMARY_ROWS:
        row = [label]
        for fy in periods:
            bs, pl = fin.bs_for(fy), fin.pl_for(fy)
            if key.startswith("pl."):
                row.append(_fmt(getattr(pl, key[3:], None) if pl else None))
            elif key == "bs._total_debt":
                ltb = val(bs.long_term_borrowings) if bs else None
                stb = val(bs.short_term_borrowings) if bs else None
                row.append("not found" if ltb is None or stb is None else f"{ltb + stb:,.0f}")
            else:
                row.append(_fmt(getattr(bs, key[3:], None) if bs else None))
        rows.append(row)
    return rows


def _ratio_table(analysis: AnalysisResult) -> list[list[str]]:
    periods = sorted(analysis.by_period)
    names: list[str] = []
    for p in periods:
        for r in analysis.by_period[p]:
            if r.name not in names:
                names.append(r.name)

    rows = [["Ratio"] + periods + ["Formula", "Data completeness"]]
    for name in names:
        results = [analysis.ratio(name, p) for p in periods]
        latest = next((r for r in reversed(results) if r), None)
        notes = {r.data_completeness.value for r in results if r}
        rows.append(
            [name]
            + [r.display if r else "n/a" for r in results]
            + [latest.formula if latest else "", ", ".join(sorted(notes))]
        )
    for r in analysis.bank:
        rows.append([f"{r.name} ({r.period})"] + [r.display] + [""] * (len(periods) - 1) + [r.formula, r.data_completeness.value])
    return rows


def _citations(fin: ExtractedFinancials) -> list[str]:
    out: list[str] = []
    for group in (fin.balance_sheets, fin.profit_and_loss, fin.bank_statements):
        for stmt in group:
            for field, value in stmt:
                if isinstance(value, Figure):
                    out.append(f"{field} = {value.value:,.0f} <- {value.citation}")
    return sorted(set(out))


def build_memo(
    *,
    fin: ExtractedFinancials,
    checklist: ChecklistReport,
    analysis: AnalysisResult,
    flags: list[RiskFlag],
    verification: list[VerificationFinding],
    narrative: str,
    trend_commentary: str,
    backend: str,
    open_questions: list[str] | None = None,
) -> CreditMemo:
    sections: list[MemoSection] = []

    sections.append(
        MemoSection(
            heading="Document checklist",
            body=(
                "Status of the documents required by the configured checklist "
                f"({checklist.checklist_name}). Items marked missing or stale must be "
                "obtained before this file is complete."
            ),
            table=[["Requirement", "Status", "Reason", "Files"]]
            + [
                [r.label, r.status.value.upper(), r.reason, ", ".join(r.matched_files) or "-"]
                for r in checklist.results
            ],
        )
    )

    sections.append(
        MemoSection(
            heading="Financial summary",
            body="All amounts in absolute rupees, as extracted from the submitted statements. "
            "'not found' means the line item was not present -- it has not been estimated.",
            table=_summary_table(fin),
        )
    )

    incomplete = [
        r
        for rs in analysis.by_period.values()
        for r in rs
        if r.data_completeness is not Completeness.COMPLETE
    ]
    sections.append(
        MemoSection(
            heading="Ratio analysis",
            body=(
                "Every ratio below is computed arithmetically from the extracted figures; "
                "the formula is shown so it can be re-derived by hand. "
                + (
                    f"{len(incomplete)} ratio-periods could not be computed and are marked "
                    "'insufficient data' rather than estimated."
                    if incomplete
                    else "All ratios were computable from the submitted documents."
                )
            ),
            table=_ratio_table(analysis),
        )
    )

    if verification:
        sections.append(
            MemoSection(
                heading="Arithmetic verification findings",
                body=(
                    "The following internal consistency checks failed on the submitted statements. "
                    "No figure has been adjusted -- these are reported for human review."
                ),
                table=[["Check", "Statement", "Expected", "As stated", "Difference", "Severity"]]
                + [
                    [
                        f.check,
                        f.statement,
                        f"{f.expected:,.0f}",
                        f"{f.actual:,.0f}",
                        f"{f.difference:,.0f}",
                        f.severity.value.upper(),
                    ]
                    for f in verification
                ],
            )
        )
    else:
        sections.append(
            MemoSection(
                heading="Arithmetic verification findings",
                body="All internal consistency checks passed on the statements that could be verified.",
            )
        )

    sections.append(
        MemoSection(
            heading="Risk flags",
            body=narrative or "No risk rules fired on the computable ratios.",
            table=[["Severity", "Flag", "Period", "Detail"]]
            + [[f.severity.value.upper(), f.label, f.period, f.message] for f in flags],
        )
    )

    sections.append(MemoSection(heading="Trend commentary", body=trend_commentary))

    questions = list(open_questions or [])
    questions += [f"Please provide: {r.label} ({r.reason})" for r in checklist.gaps]
    questions += [
        f"Please clarify: {f.check} on {f.statement} -- stated {f.actual:,.0f} vs {f.expected:,.0f} expected."
        for f in verification
    ]
    for r in incomplete:
        questions.append(
            f"{r.name} ({r.period}) could not be computed -- missing {', '.join(r.missing_inputs) or 'inputs'}."
        )

    return CreditMemo(
        borrower_name=fin.borrower_name or "Borrower name not established",
        facility_requested=fin.facility_requested or "Facility not stated in the submitted documents",
        prepared_on=date.today(),
        backend=backend,
        sections=sections,
        citations=_citations(fin),
        open_questions=questions,
    )


def write_docx(memo: CreditMemo, path: Path) -> Path:
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    doc.add_heading("Credit Memo (Draft)", level=0)

    warning = doc.add_paragraph()
    run = warning.add_run(DISCLAIMER)
    run.bold = True
    run.font.size = Pt(10)

    header = doc.add_table(rows=0, cols=2)
    header.style = "Table Grid"
    for key, value in [
        ("Borrower", memo.borrower_name),
        ("Facility requested", memo.facility_requested),
        ("Prepared on", memo.prepared_on.isoformat()),
        ("LLM backend", f"{memo.backend}" + (" (local, offline)" if memo.backend == "ollama" else " (external API)")),
    ]:
        cells = header.add_row().cells
        cells[0].text = key
        cells[1].text = str(value)

    for section in memo.sections:
        doc.add_heading(section.heading, level=1)
        if section.body:
            doc.add_paragraph(section.body)
        if len(section.table) > 1:
            table = doc.add_table(rows=1, cols=len(section.table[0]))
            table.style = "Table Grid"
            for i, head in enumerate(section.table[0]):
                table.rows[0].cells[i].text = str(head)
            for row in section.table[1:]:
                cells = table.add_row().cells
                for i, value in enumerate(row[: len(section.table[0])]):
                    cells[i].text = str(value)
        elif section.table:
            doc.add_paragraph("No entries.")

    doc.add_heading("Open questions for the borrower", level=1)
    if memo.open_questions:
        for q in memo.open_questions:
            doc.add_paragraph(q, style="List Bullet")
    else:
        doc.add_paragraph("None arising from the automated review.")

    doc.add_heading("Source citations", level=1)
    doc.add_paragraph("Every figure used above, with the document, page/sheet and row it came from.")
    for c in memo.citations:
        doc.add_paragraph(c, style="List Bullet")

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))
    return path
